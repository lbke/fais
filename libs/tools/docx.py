from docx import Document

from langchain.tools import ToolRuntime, tool
from os import path


@tool
def create_empty_docx_document(runtime: ToolRuntime, filepath: str) -> str:
    """
    Initialize an empty docx document
    Use to create new documents as requested by the user
    Do NOT use to edit existing documents or adapt a template, only to create new documents
    """
    if path.exists(filepath):
        return f"Error: File {filepath} already exists."
    doc = Document().save(filepath)
    runtime.context[filepath] = doc
    return f"Empty document created at {filepath}"


@tool
def add_heading_docx(runtime: ToolRuntime, filepath: str, heading: str, level: int = 0):
    """
    Add a heading to the document.
    #docx.document.Document.add_heading
    Do NOT use to edit existing documents or adapt a template, only to create new documents as requested by the user.
    Documentation of add_heading function in python-docx : https://python-docx.readthedocs.io/en/latest/api/document.html
    Args:
        level: 0 is the main Title, 1 is heading of level 1 etc.
    """
    doc = runtime.context[filepath]
    if not doc:
        return f"Error: document {filepath} doesn't exist. Call {create_empty_docx_document.__repr_name__} first."
    doc.add_heading(heading, level=level)
    doc.save()


@tool
def add_paragraph_docx(runtime: ToolRuntime, filepath: str, text: str):
    """
    Add a paragraph to the document.
    Do NOT use to edit existing documents or adapt a template, only to create new documents as requested by the user.
    Documentation of add_paragraph function in python-docx : https://python-docx.readthedocs.io/en/latest/api/document.html#docx.document.Document.add_paragraph
    Args:
        text: The text of the paragraph to add.
    """
    doc = runtime.context[filepath]
    if not doc:
        return f"Error: document {filepath} doesn't exist. Call {create_empty_docx_document.__repr_name__} first."
    doc.add_paragraph(text)
    doc.save()


TOOLS = [create_empty_docx_document, add_heading_docx, add_paragraph_docx]
TOOLS_DESCRIPTION = "Creating simple docx documents from scratch"
